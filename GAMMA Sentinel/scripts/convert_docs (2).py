# -*- coding: utf-8 -*-
import os
import sys
import subprocess

def install_and_import(package, import_name=None):
    if import_name is None:
        import_name = package
    try:
        __import__(import_name)
    except ImportError:
        print(f"Instalando dependencia necesaria: {package}...")
        try:
            subprocess.check_call([sys.executable, "-m", "pip", "install", package])
        except Exception as e:
            print(f"Error al instalar {package}: {e}")
            print(f"Por favor, instala la librería manualmente ejecutando: pip install {package}")
            sys.exit(1)

# Asegurar dependencias de lectura
install_and_import("pypdf")
install_and_import("python-docx", "docx")
install_and_import("openpyxl")

import pypdf
import docx
import openpyxl

# Directorio de los archivos
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
FILES_DIR = os.path.join(SCRIPT_DIR, "..", "knowledge", "files")

# Mapeo de archivos originales a los nombres de salida .txt esperados por el frontend
MAPPINGS = {
    "Base_Conocimiento_GAMMA.txt": "Base_Conocimiento_GAMMA.txt",
    "Archivo_SharePoint_GAMMA.xlsx": "Archivo_SharePoint_GAMMA.txt",
    "Backlog Gamma Mantenimiento.jmg.xlsx": "BacklogGammaMantenimiento.txt",
    "Backlog Gamma Mantenimiento.xlsx": "BacklogGammaMantenimiento.txt",
    "GAMMA - Manual de Usuario v1.0.pdf": "GAMMA-ManualdeUsuariov1.0.txt",
    "GAMMA - Manual de Usuario v1.2.pdf": "GAMMA-ManualdeUsuariov1.0.txt",
    "Lote Movil - Manual de Usuario v1.2.pdf": "LoteMovil-ManualdeUsuariov1.2.txt",
    "LoteMovil - Manual de Usuario v1.2.pdf": "LoteMovil-ManualdeUsuariov1.2.txt",
    "LoteMovil - Manual de Administrador v1.2.pdf": "LoteMovil-ManualdeAdministradorv1.2.txt",
    "Manual_RT_GAMMA_Integrado.docx": "Manual_RT_GAMMA_Integrado.txt",
    "NOA - Manual de Usuario v1.2.pdf": "NOA-ManualdeUsuariov1.2.txt"
}

def clean_text(text):
    if not text:
        return ""
    # Normalizar retornos de carro y espacios
    return "\n".join(line.strip() for line in text.splitlines() if line.strip())

def extract_pdf(file_path):
    print(f"Leyendo PDF: {os.path.basename(file_path)}...")
    text_content = []
    with open(file_path, "rb") as f:
        reader = pypdf.PdfReader(f)
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                text_content.append(f"--- PÁGINA {i+1} ---\n{text}")
    return "\n\n".join(text_content)

def extract_docx(file_path):
    print(f"Leyendo DOCX: {os.path.basename(file_path)}...")
    doc = docx.Document(file_path)
    text_content = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_content.append(paragraph.text.strip())
    # Extraer también de tablas si existen
    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_data:
                text_content.append(" | ".join(row_data))
    return "\n\n".join(text_content)

def extract_xlsx(file_path):
    print(f"Leyendo XLSX: {os.path.basename(file_path)}...")
    wb = openpyxl.load_workbook(file_path, data_only=True)
    text_content = []
    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        text_content.append(f"=== HOJA: {sheet_name} ===")
        for row in sheet.iter_rows(values_only=True):
            row_data = [str(cell).replace('\r', ' ').replace('\n', ' ').strip() for cell in row if cell is not None]
            if row_data:
                text_content.append(" | ".join(row_data))
    return "\n\n".join(text_content)

def main():
    if not os.path.exists(FILES_DIR):
        print(f"Error: No se encontró la carpeta de archivos en {FILES_DIR}")
        return

    print("=" * 60)
    print("Iniciando conversión de base de conocimientos a texto plano...")
    print(f"Carpeta origen/destino: {os.path.abspath(FILES_DIR)}")
    print("=" * 60)

    converted_count = 0
    skipped_count = 0

    # Escanear archivos en la carpeta
    available_files = os.listdir(FILES_DIR)
    
    for filename in available_files:
        # Ignorar archivos temporales o de metadatos .mcs.yml
        if filename.startswith("~$") or filename.endswith(".mcs.yml"):
            continue

        matched_key = None
        # Buscar coincidencia exacta o coincidencia por nombre normalizado
        for original_name in MAPPINGS.keys():
            if filename.lower() == original_name.lower():
                matched_key = original_name
                break
        
        if not matched_key:
            # Si no hay coincidencia exacta de mapeo, pero es un archivo soportado, 
            # avisar al usuario para que lo renombre
            ext = os.path.splitext(filename)[1].lower()
            if ext in [".pdf", ".docx", ".xlsx", ".txt"]:
                print(f"[Aviso] Se encontró '{filename}' pero su nombre no coincide con las fuentes mapeadas del frontend.")
                print(f"        Si deseas indexar este archivo, renómbralo para que coincida exactamente con alguno de estos:")
                for name in MAPPINGS.keys():
                    print(f"        - {name}")
            continue

        source_path = os.path.join(FILES_DIR, filename)
        dest_filename = MAPPINGS[matched_key]
        dest_path = os.path.join(FILES_DIR, dest_filename)

        # Si el origen ya es un archivo .txt, solo validamos que exista, no requiere conversión
        if filename.lower().endswith(".txt"):
            print(f"[OK] '{filename}' ya está en formato texto plano (.txt).")
            continue

        try:
            ext = os.path.splitext(filename)[1].lower()
            extracted_text = ""
            
            if ext == ".pdf":
                extracted_text = extract_pdf(source_path)
            elif ext == ".docx":
                extracted_text = extract_docx(source_path)
            elif ext == ".xlsx":
                extracted_text = extract_xlsx(source_path)

            if extracted_text.strip():
                cleaned = clean_text(extracted_text)
                with open(dest_path, "w", encoding="utf-8") as f:
                    f.write(cleaned)
                print(f"-> Convertido con éxito a: {dest_filename}")
                converted_count += 1
            else:
                print(f"[Advertencia] No se pudo extraer texto de {filename} (¿está vacío?)")
        except Exception as e:
            print(f"Error procesando {filename}: {e}")

    print("=" * 60)
    print(f"Proceso finalizado. Conversiones exitosas: {converted_count}")
    print("¡Listo! Si agregaste nuevos archivos, no olvides subirlos a tu repositorio con Git.")
    print("=" * 60)

if __name__ == "__main__":
    main()
